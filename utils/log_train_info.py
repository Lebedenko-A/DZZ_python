import numpy as np
from docx import Document
from docx.shared import Inches
from datetime import datetime
from matplotlib import pyplot as plt
import csv

def parse_filt_noise_window(str):
    return str.split(" ")

def barplot(data, labels = ["CPU", "GPU", "Python"]):
    width = 1.5
    x_cord = np.array([3, 5, 7, 9])
    for key in data:
        plt.figure()
        for i, window in enumerate(data[key]):
            plt.bar(x_cord[i] - width/3, data[key][window][0], width/3, label="CPU", color="blue")
            plt.bar(x_cord[i], data[key][window][1], width/3, label="GPU", color="green")
            plt.bar(x_cord[i] + width / 3, data[key][window][2], width/3, label="Python", color="red")
            #plt.bar(data[key][window][0])
        plt.legend(loc=0)
        plt.xlabel("Window size")
        plt.ylabel("Time, s")
        plt.title(str(key))
        plt.savefig("bar" + str(key) + ".png")
        plt.show()


def load_csv_to_dict(filename):
    result = dict()
    with open(filename, newline="") as csvfile:
        reader = csv.reader(csvfile, delimiter=' ', quotechar="|")
        for i, row in enumerate(reader):
            if i != 0:
                im_info = row[0].split(",")
                print(im_info)
                result[im_info[0] + " " + im_info[1] + " " + im_info[2]] = im_info[3:]
    return result

def create_docx(result_log, cpu_filename, gpu_filename, filename=""):
    timenow = datetime.now().strftime("%d-%m-%Y_at_%H:%M")

    document = Document()

    log_gpu = load_csv_to_dict(gpu_filename)
    log_cpu = load_csv_to_dict(cpu_filename)

    document.add_heading("Metric Log Image " + filename, 0)

    logdata_history = []

    table_top = ["Filter name", "Window size", "CPU (C++) TIME", "GPU (CUDA) TIME", "Python (numpy) TIME"]
    table = document.add_table(rows=1, cols=len(table_top))
    for i, t in enumerate(table_top):
        table.rows[0].cells[i].text = t

    '''table_top_2 = ["", "", "", "MSE", "PSNR", "PSNRHVS", "PSNRHVSM", "TIME", "MSE", "PSNR", "PSNRHVS", "PSNRHVSM", "TIME", "MSE", "PSNR", "PSNRHVS", "PSNRHVSM", "TIME"]
    table = document.add_table(rows=1, cols=len(table_top_2))
    for i, t in enumerate(table_top_2):
        table.rows[0].cells[i].text = t
'''
    data4plot = dict()
    cout_of_row = 0
    for key_cpu, key_gpu, key_python in zip(log_cpu, log_gpu, result_log):
        first_3 = parse_filt_noise_window(key_cpu)
        if data4plot.get(first_3[0]) == None: data4plot[first_3[0]] = dict()
        if data4plot[first_3[0]].get(first_3[2]) == None: data4plot[first_3[0]][first_3[2]] = []
        row = table.add_row().cells

        row[0].text = first_3[0]
        row[1].text = first_3[2]
        row[2].text = log_cpu[key_cpu][-1]
        data4plot[first_3[0]][first_3[2]].append(float(log_cpu[key_cpu][-1]))
        row[3].text = log_gpu[key_gpu][-1]
        data4plot[first_3[0]][first_3[2]].append(float(log_gpu[key_gpu][-1]))
        row[4].text = result_log[key_python][-1]
        data4plot[first_3[0]][first_3[2]].append(float(result_log[key_python][-1]))

        cout_of_row += 1
        if cout_of_row == 13: break
    barplot(data4plot)
    document.add_picture('barMedian.png', height=Inches(4))
    document.add_picture('barLi.png', height=Inches(4))
    document.add_picture('barFrost.png', height=Inches(4))
    document.add_picture('barDCT_based.png', height=Inches(4))
    document.save("test" + timenow + ".docx")